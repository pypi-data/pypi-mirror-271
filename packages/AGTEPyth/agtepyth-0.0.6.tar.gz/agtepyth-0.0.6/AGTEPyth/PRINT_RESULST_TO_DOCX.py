from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Kiểm tra cục bộ: Với các loại hàm nhập trung gian như thế này thì cần phải có giá trị chạy thử cho chúng để kiểm tra trước
'''T3_pol=1111
Pi_polnoe=[2,4]
MH=1
Sigma_vx=1
KPD_VnutT=1
k_ispol=1
TH=111
Sigma_kom=1
h0_cr=111
Ne=11111
DeltaT_Kom=1
Sigma_ks=1
g_otb_=1
X=1
DeltaT_VozOxl=1
Fi_c=1
g_utech_=1
KPD_vnTV=1
T_ct=1111
KPD_ks=1
alpha_=1
KPD_reduk=1
ColumnPik=[1,2]
ColumnT2pol=[1,2]
ColumnC_pmiB=[1,2]
ColumnTg=[1,2]
ColumnTq=[1,2]
ColumnT4a=[1,2]
ColumnT4q=[1,2]
ColumnC_pmig=[1,2]
ColumnT5t=[1,2]
ColumnP3naP5=[1,2]
ColumnHor=[1,2]
ColumnHog=[1,2]
Columnalpha=[1,2]
ColumnZ=[1,2]
Columnqoxl=[1,2]
ColumnPioxl=[1,2]
ColumnHoxl_=[1,2]
ColumnKAPq=[1,2]
ColumnKAPoxl=[1,2]
ColumnKPDad=[1,2]
Columnlcv=[1,2]
Columngoxl_=[1,2]
ColumnC5_=[1,2]
ColumnGv=[1,2]
ColumnGg=[1,2]
ColumnNeud=[1,2]
ColumnCe=[1,2]
TipKom='однокаскадный осевой компрессор'
UserName='Phạm Thành Quyết'
UserGroup='6020407/20006'
folderNameSelect='''

def InFileResult(T3_pol,Pi_polnoe,MH,Sigma_vx,KPD_VnutT,k_ispol,TH,Sigma_kom,h0_cr,Ne,DeltaT_Kom,Sigma_ks,g_otb_,X,DeltaT_VozOxl,Fi_c,g_utech_,
                KPD_vnTV,T_ct,KPD_ks,alpha_,KPD_reduk,ColumnPik,ColumnT2pol,ColumnC_pmiB,ColumnTg,ColumnTq,ColumnT4a,ColumnT4q,ColumnC_pmig,ColumnT5t,
                ColumnP3naP5,ColumnHor,ColumnHog,Columnalpha,ColumnZ,Columnqoxl,ColumnPioxl,ColumnHoxl_,ColumnKAPq,ColumnKAPoxl,ColumnKPDad,Columnlcv,
                Columngoxl_,ColumnC5_,ColumnGv,ColumnGg,ColumnNeud,ColumnCe,TipKom,UserName,UserGroup,folderNameSelect):
    document = Document()

    DauMuc = document.add_paragraph() #Tạo đoạn văn để thêm đầu đề
    DauMuc.alignment = WD_ALIGN_PARAGRAPH.CENTER #Căn giữa
    DauMuc_=DauMuc.add_run()
    DauMuc_.bold = True
    DauMuc_.font.name = 'Times New Roman'
    DauMuc_.font.size=Pt(12)
    DauMuc_.add_text('''ПPOГPAММA "A2GTPyt"''') #In đậm

    MoTaPM = document.add_paragraph() #Tạo đoạn văn để thêm mô tả phần mềm
    MoTaPM.alignment = WD_ALIGN_PARAGRAPH.CENTER
    MoTaPM_=MoTaPM.add_run()
    MoTaPM_.bold = True
    MoTaPM_.font.name = 'Times New Roman'
    MoTaPM_.font.size=Pt(12)
    MoTaPM_.add_text('Вapиaнтный pacчeт пapaмeтpoв paбoчeгo пpoцecca aвиaциoннoгo турбовального двигателя c oxлaждaeмoй тypбинoй')

    HoTen = document.add_paragraph() #Tạo đoạn văn để thêm đầu đề
    HoTen.alignment = WD_ALIGN_PARAGRAPH.LEFT
    HoTen_=HoTen.add_run()
    HoTen_.bold = True
    HoTen_.font.name = 'Times New Roman'
    HoTen_.font.size=Pt(12)
    HoTen_.add_text('Фамилия и имя: ')
    HoTen_ad=HoTen.add_run()
    HoTen_ad.font.name = 'Times New Roman'
    HoTen_ad.font.size=Pt(12)
    HoTen_ad.add_text(UserName)

    Nhom = document.add_paragraph() #Tạo đoạn văn để thêm đầu đề
    Nhom.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    Nhom_=Nhom.add_run()
    Nhom_.bold = True
    Nhom_.font.name = 'Times New Roman'
    Nhom_.font.size=Pt(12)
    Nhom_.add_text('Группа: ')
    Nhom_ad=Nhom.add_run()
    Nhom_ad.font.name = 'Times New Roman'
    Nhom_ad.font.size=Pt(12)
    Nhom_ad.add_text(UserGroup)

    Datavx = document.add_paragraph() #Tạo đoạn văn để thêm đầu đề
    Datavx.alignment = WD_ALIGN_PARAGRAPH.CENTER #Căn giữa
    Datavx_=Datavx.add_run()
    Datavx_.font.name = 'Times New Roman'
    Datavx_.font.size=Pt(12)
    Datavx_.font.underline = True #Thêm gạch dưới
    Datavx_.add_text('ИСХОДНЫЕ ДАННЫЕ:')

    T3text = document.add_paragraph() #Tạo đoạn văn để thêm đầu đề
    T3text.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    T3text_=T3text.add_run()
    T3text_.font.name = 'Times New Roman'
    T3text_.font.size=Pt(12)
    T3text_.bold=True
    T3text_.add_text('Т3* = ')
    T3text_ad=T3text.add_run()
    T3text_ad.font.name = 'Times New Roman'
    T3text_ad.font.size=Pt(12)
    T3text_ad.add_text(str(T3_pol)+' K')

    Pik = document.add_paragraph() #Tạo đoạn văn để thêm đầu đề
    Pik.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    Pik_=Pik.add_run()
    Pik_.bold=True
    Pik_.font.name = 'Times New Roman'
    Pik_.font.size=Pt(12)
    Pik_.add_text('ПИК = ')
    Pik_ad=Pik.add_run()
    Pik_ad.font.name = 'Times New Roman'
    Pik_ad.font.size=Pt(12)
    Pik_ad.add_text(', '.join([str(i) for i in Pi_polnoe]))

    TipK = document.add_paragraph() #Tạo đoạn văn để thêm đầu đề
    TipK.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    TipK_=TipK.add_run()
    TipK_.font.name = 'Times New Roman'
    TipK_.bold=True
    TipK_.font.size=Pt(12)
    TipK_.add_text('Тип компрессора: ')
    TipK_ad=TipK.add_run()
    TipK_ad.font.name = 'Times New Roman'
    TipK_ad.font.size=Pt(12)
    TipK_ad.add_text(TipKom)

    #TẠO BẢNG DỮ LIỆU NHẬP VÀO:

    DataInput = (('Тн (К)',TH,'СИГком',Sigma_kom,'hocp (кДж/кг)',h0_cr,'Ne (кВт)',Ne),
        ('Дтком (К)',DeltaT_Kom,'СИГкс',Sigma_ks,"g’_oтб",g_otb_,'Х',X),
        ('ДТв (К)',DeltaT_VozOxl,'ПИс',Fi_c,"g’_ут",g_utech_,'КПДтв',KPD_vnTV),
        ('Тст (К)',T_ct,'КПДкс',KPD_ks,"Альфа*",alpha_,'КПДред',KPD_reduk)) #Tạo 1 đối tượng tuple kiểu ma trận chứa các giá trị để thêm vào bảng
    #Trong trường hợp này không tạo hàng đầu tên cột nên cần bỏ 1 hàng giá trị trên cùng trong tubple và điền trực tiếp hàng giá trị đó vào hàng đầu của bảng

    tableDataInput = document.add_table(rows=1, cols=8) #Tạo bảng ban đầu 1 hàng 3 cột
    tableDataInput.style = "Table Grid"
    Input_cells = tableDataInput.rows[0].cells #Tạo hàng đầu để điền trực tiếp các giá trị đầu (với bảng không có header)
    Input_cells[0].text = 'Мн'
    Input_cells[1].text = str(MH)
    Input_cells[2].text = 'СИГвх'
    Input_cells[3].text = str(Sigma_vx)
    Input_cells[4].text = 'КПДвт'
    Input_cells[5].text = str(KPD_VnutT)
    Input_cells[6].text = 'К_исп'
    Input_cells[7].text = str(k_ispol)
    for a1, a2, a3, a4, a5, a6, a7, a8 in DataInput:
        Inp_cells = tableDataInput.add_row().cells
        Inp_cells[0].text = a1
        Inp_cells[1].text = str(a2)
        Inp_cells[2].text = a3
        Inp_cells[3].text = str(a4)
        Inp_cells[4].text = a5
        Inp_cells[5].text = str(a6)
        Inp_cells[6].text = a7
        Inp_cells[7].text = str(a8)
    for row in tableDataInput.rows: #Vì các định dạng văn bản chỉ có trong đối tượng paragraph nên cần phải tạo vòng lặp nếu cần định dạng tất cả các cells như nhau.
        for cell in row.cells:
            cp = cell.paragraphs[0].runs
            cp[0].font.name = 'Time New Roman'
            cp[0].font.size=Pt(10)
    for cell in tableDataInput.columns[0].cells:
        cell.width = Inches(1)
        cp = cell.paragraphs[0].runs
        cp[0].font.bold=True
    for cell in tableDataInput.columns[1].cells:
        cell.width = Inches(0.7)
    for cell in tableDataInput.columns[2].cells:
        cell.width = Inches(0.9)
        cp = cell.paragraphs[0].runs
        cp[0].font.bold=True
    for cell in tableDataInput.columns[3].cells:
        cell.width = Inches(0.7)
    for cell in tableDataInput.columns[4].cells:
        cell.width = Inches(1.4)
        cp = cell.paragraphs[0].runs
        cp[0].font.bold=True
    for cell in tableDataInput.columns[5].cells:
        cell.width = Inches(0.7)
    for cell in tableDataInput.columns[6].cells:
        cell.width = Inches(0.9)
        cp = cell.paragraphs[0].runs
        cp[0].font.bold=True
    for cell in tableDataInput.columns[7].cells:
        cell.width = Inches(0.7)

    #TẠO PHẦN IN KẾT QUẢ
    Result = document.add_paragraph('\n') #Tạo đoạn văn để thêm đầu đề
    Result.alignment = WD_ALIGN_PARAGRAPH.CENTER #Căn giữa
    Result_=Result.add_run()
    Result_.font.name = 'Times New Roman'
    Result_.font.size=Pt(12)
    Result_.font.underline = True #Thêm gạch dưới
    Result_.add_text('РЕЗУЛЬТАТЫ:')

    DataResult1 =tuple((ColumnPik[i],ColumnT2pol[i],ColumnC_pmiB[i],ColumnTg[i],ColumnTq[i],ColumnT4a[i],ColumnT4q[i],ColumnC_pmig[i],ColumnT5t[i],ColumnP3naP5[i]) for i in range(len(Pi_polnoe)))

    tableResult1 = document.add_table(rows=1, cols=10) #Tạo bảng ban đầu 1 hàng 3 cột
    tableResult1.style = "Table Grid"
    Result_cells1 = tableResult1.rows[0].cells #Tạo hàng đầu để điền trực tiếp các giá trị đầu (với bảng không có header)
    Result_cells1[0].text = 'ПИК'
    Result_cells1[1].text = 'Т2*'
    Result_cells1[2].text = 'С_рмив'
    Result_cells1[3].text = 'Тg'
    Result_cells1[4].text = 'Tq'
    Result_cells1[5].text = 'Т4а*'
    Result_cells1[6].text = 'Т4q*'
    Result_cells1[7].text = 'С_рмиг'
    Result_cells1[8].text = 'Т5t'
    Result_cells1[9].text = 'Р3*/Р5'
    for a1, a2, a3, a4, a5, a6, a7, a8, a9, a10 in DataResult1:
        Result_cells1 = tableResult1.add_row().cells
        Result_cells1[0].text = str(a1)
        Result_cells1[1].text = str(a2)
        Result_cells1[2].text = str(a3)
        Result_cells1[3].text = str(a4)
        Result_cells1[4].text = str(a5)
        Result_cells1[5].text = str(a6)
        Result_cells1[6].text = str(a7)
        Result_cells1[7].text = str(a8)
        Result_cells1[8].text = str(a9)
        Result_cells1[9].text = str(a10)
    for row in tableResult1.rows: #Vì các định dạng văn bản chỉ có trong đối tượng paragraph nên cần phải tạo vòng lặp nếu cần định dạng tất cả các cells như nhau.
        for cell in row.cells:
            cp = cell.paragraphs[0].runs
            cp[0].font.name = 'Time New Roman'
            cp[0].font.size=Pt(10)
    for cell in tableResult1.rows[0].cells:
        cp = cell.paragraphs[0].runs
        cp[0].font.bold=True
    for cell in tableResult1.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in tableResult1.columns[1].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[2].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[3].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[4].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[5].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[6].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[7].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[8].cells:
        cell.width = Inches(0.7)
    for cell in tableResult1.columns[9].cells:
        cell.width = Inches(0.7)

    XuongDong1 = document.add_paragraph('')

    DataResult2 =tuple((ColumnPik[i],ColumnHor[i],ColumnHog[i],Columnalpha[i],ColumnZ[i],Columnqoxl[i],ColumnPioxl[i],ColumnHoxl_[i],ColumnKAPq[i],ColumnKAPoxl[i]) for i in range(len(Pi_polnoe)))

    tableResult2 = document.add_table(rows=1, cols=10) #Tạo bảng ban đầu 1 hàng 3 cột
    tableResult2.style = "Table Grid"
    Result_cells2 = tableResult2.rows[0].cells #Tạo hàng đầu để điền trực tiếp các giá trị đầu (với bảng không có header)
    Result_cells2[0].text = 'ПИК'
    Result_cells2[1].text = 'Нор'
    Result_cells2[2].text = 'Ног*'
    Result_cells2[3].text = 'Альфа'
    Result_cells2[4].text = 'Z'
    Result_cells2[5].text = 'q_охл'
    Result_cells2[6].text = 'ПИ_охл'
    Result_cells2[7].text = "Н’_охл"
    Result_cells2[8].text = 'КАП_q'
    Result_cells2[9].text = 'КАП_охл'
    for a1, a2, a3, a4, a5, a6, a7, a8, a9, a10 in DataResult2:
        Result_cells2 = tableResult2.add_row().cells
        Result_cells2[0].text = str(a1)
        Result_cells2[1].text = str(a2)
        Result_cells2[2].text = str(a3)
        Result_cells2[3].text = str(a4)
        Result_cells2[4].text = str(a5)
        Result_cells2[5].text = str(a6)
        Result_cells2[6].text = str(a7)
        Result_cells2[7].text = str(a8)
        Result_cells2[8].text = str(a9)
        Result_cells2[9].text = str(a10)
    for row in tableResult2.rows: #Vì các định dạng văn bản chỉ có trong đối tượng paragraph nên cần phải tạo vòng lặp nếu cần định dạng tất cả các cells như nhau.
        for cell in row.cells:
            cp = cell.paragraphs[0].runs
            cp[0].font.name = 'Time New Roman'
            cp[0].font.size=Pt(10)
    for cell in tableResult2.rows[0].cells:
        cp = cell.paragraphs[0].runs
        cp[0].font.bold=True
    for cell in tableResult2.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in tableResult2.columns[1].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[2].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[3].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[4].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[5].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[6].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[7].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[8].cells:
        cell.width = Inches(0.7)
    for cell in tableResult2.columns[9].cells:
        cell.width = Inches(0.7)

    XuongDong2 = document.add_paragraph('')

    DataResult3 =tuple((ColumnPik[i],ColumnKPDad[i],Columnlcv[i],Columngoxl_[i],ColumnC5_[i],ColumnGv[i],ColumnGg[i],ColumnNeud[i],ColumnCe[i]) for i in range(len(Pi_polnoe)))

    tableResult3 = document.add_table(rows=1, cols=9) #Tạo bảng ban đầu 1 hàng 3 cột
    tableResult3.style = "Table Grid"
    Result_cells3 = tableResult3.rows[0].cells #Tạo hàng đầu để điền trực tiếp các giá trị đầu (với bảng không có header)
    Result_cells3[0].text = 'ПИК'
    Result_cells3[1].text = 'КПДад'
    Result_cells3[2].text = "l'_св"
    Result_cells3[3].text = "g'_ охл"
    Result_cells3[4].text = 'C5'
    Result_cells3[5].text = 'Gв'
    Result_cells3[6].text = 'Gг'
    Result_cells3[7].text = 'Nеуд'
    Result_cells3[8].text = 'Се'
    for a1, a2, a3, a4, a5, a6, a7, a8, a9 in DataResult3:
        Result_cells3 = tableResult3.add_row().cells
        Result_cells3[0].text = str(a1)
        Result_cells3[1].text = str(a2)
        Result_cells3[2].text = str(a3)
        Result_cells3[3].text = str(a4)
        Result_cells3[4].text = str(a5)
        Result_cells3[5].text = str(a6)
        Result_cells3[6].text = str(a7)
        Result_cells3[7].text = str(a8)
        Result_cells3[8].text = str(a9)
    for row in tableResult3.rows: #Vì các định dạng văn bản chỉ có trong đối tượng paragraph nên cần phải tạo vòng lặp nếu cần định dạng tất cả các cells như nhau.
        for cell in row.cells:
            cp = cell.paragraphs[0].runs
            cp[0].font.name = 'Time New Roman'
            cp[0].font.size=Pt(10)
    for cell in tableResult3.rows[0].cells:
        cp = cell.paragraphs[0].runs
        cp[0].font.bold=True
    for cell in tableResult3.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in tableResult3.columns[1].cells:
        cell.width = Inches(0.8)
    for cell in tableResult3.columns[2].cells:
        cell.width = Inches(0.8)
    for cell in tableResult3.columns[3].cells:
        cell.width = Inches(0.8)
    for cell in tableResult3.columns[4].cells:
        cell.width = Inches(0.8)
    for cell in tableResult3.columns[5].cells:
        cell.width = Inches(0.8)
    for cell in tableResult3.columns[6].cells:
        cell.width = Inches(0.8)
    for cell in tableResult3.columns[7].cells:
        cell.width = Inches(0.8)
    for cell in tableResult3.columns[8].cells:
        cell.width = Inches(0.8)

    result='Result'+str(T3_pol)+'K.docx'
    try:
        document.save(folderNameSelect+'\\'+result)
        #document.save(result) #Kiểm tra cục bộ
    except PermissionError:
        i=1
        while 2>1:
            try:
                result_=result.replace('K.docx','(%s)'%str(i)+'K.docx')
                document.save(folderNameSelect+'\\'+result_)
                #document.save(result_) #Kiểm tra cục bộ
                break
            except:
                i+=1
'''InFileResult(T3_pol,Pi_polnoe,MH,Sigma_vx,KPD_VnutT,k_ispol,TH,Sigma_kom,h0_cr,Ne,DeltaT_Kom,Sigma_ks,g_otb_,X,DeltaT_VozOxl,Fi_c,g_utech_,
                KPD_vnTV,T_ct,KPD_ks,alpha_,KPD_reduk,ColumnPik,ColumnT2pol,ColumnC_pmiB,ColumnTg,ColumnTq,ColumnT4a,ColumnT4q,ColumnC_pmig,ColumnT5t,
                ColumnP3naP5,ColumnHor,ColumnHog,Columnalpha,ColumnZ,Columnqoxl,ColumnPioxl,ColumnHoxl_,ColumnKAPq,ColumnKAPoxl,ColumnKPDad,Columnlcv,
                Columngoxl_,ColumnC5_,ColumnGv,ColumnGg,ColumnNeud,ColumnCe,TipKom,UserName,UserGroup,folderNameSelect)''' #Kiểm tra cục bộ